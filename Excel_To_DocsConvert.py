import pandas as pd
from docx import Document

# Read the Excel file
df = pd.read_excel('Startech.xlsx')

# Create a new Word document
doc = Document()

# Iterate through the rows and write questions and answers to the document
for index, row in df.iterrows():
    question = f'Q{index+1}) {row["question"]}'
    answer = f'A{index+1}) {row["answer"]}'
    doc.add_paragraph(question)
    doc.add_paragraph(answer)
    doc.add_paragraph("")  # Add an empty paragraph for separation

# Save the document
doc.save('startech_qa.docx')

print("Word document created successfully!")